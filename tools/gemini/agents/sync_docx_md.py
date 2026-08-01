# -*- coding: utf-8 -*-
"""sync_docx_md — sincroniza para o .md as edições que o humano fez no .docx.

O .md é a FONTE DA VERDADE. Quando o humano edita o Word (mais fácil para ele),
este script traz as alterações de volta para o .md, para que a próxima geração
do .docx não as destrua.

Uso:
    python sync_docx_md.py <editado.docx> <destino.md>            # só relata
    python sync_docx_md.py <editado.docx> <destino.md> --apply    # relata e aplica

Como funciona (não improvise outro método — este já está validado):
  1. Gera, a partir do .md atual, um .docx "esperado" (via md2docx).
  2. Compara esperado x editado usando a MESMA extração de texto nos dois.
     Comparar docx-vs-docx dá ~97% de precisão; comparar .md-vs-.docx
     direto dá ~34% (granularidades diferentes) e gera ruído inútil.
  3. Relata só os blocos que mudaram, com o número da linha no .md.
  4. Com --apply, aplica sozinho as substituições de texto seguras
     (1:1, texto puro, ocorrência única) e deixa o resto para decisão humana.

Limite conhecido: formatação feita no Word (fonte, cor, largura de coluna,
tamanho de imagem) não existe em Markdown e não é sincronizável.
Texto, células de tabela, linhas novas e imagens inseridas sincronizam bem.
"""
import sys
import os
import re
import difflib
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import md2docx  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.table import Table  # noqa: E402


# ---------------- extração ----------------
def _text_of(el):
    """Todo o texto do elemento, inclusive dentro de w:hyperlink."""
    return ''.join(t.text or '' for t in el.iter(qn('w:t')))


def extract_blocks(path):
    """Um bloco por parágrafo e por célula (células mescladas contam uma vez).

    ATENÇÃO: a deduplicação guarda os próprios elementos em `vivos`. Usar apenas
    `id(cell._tc)` é um bug — o lxml libera o proxy e reaproveita o id, fazendo o
    script pular blocos legítimos (já causou falso negativo em edição real).
    """
    doc = Document(path)
    seen, vivos, out = set(), [], []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            t = _text_of(child).strip()
            if t:
                out.append(t)
        elif child.tag == qn('w:tbl'):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    tc = cell._tc
                    vivos.append(tc)          # mantém vivo: impede reuso de id
                    if id(tc) in seen:
                        continue
                    seen.add(id(tc))
                    t = _text_of(tc).strip()
                    if t:
                        out.append(t)
    return out


def count_images(path):
    return len(Document(path).element.body.findall('.//' + qn('w:drawing')))


def extrair_imagens(docx_path, md_path):
    """Extrai as imagens do .docx para <pasta do md>/images/.

    Informa, para cada uma, o último texto que a precede — é o que permite
    referenciá-la na seção certa do .md.
    """
    doc = Document(docx_path)
    destino = os.path.join(os.path.dirname(md_path), 'images')
    os.makedirs(destino, exist_ok=True)
    base = os.path.splitext(os.path.basename(md_path))[0]
    base = re.sub(r'^documento-analise-', '', base)
    ultimo, achadas, i = '(topo)', [], 0
    for child in doc.element.body.iterchildren():
        if child.tag != qn('w:p'):
            continue
        for blip in child.findall('.//' + qn('a:blip')):
            rid = blip.get(qn('r:embed'))
            part = doc.part.related_parts[rid]
            ext = os.path.splitext(str(part.partname))[1] or '.png'
            i += 1
            nome = '%s-img%02d%s' % (base, i, ext)
            with open(os.path.join(destino, nome), 'wb') as f:
                f.write(part.blob)
            achadas.append((nome, ultimo))
        t = ''.join(x.text or '' for x in child.iter(qn('w:t'))).strip()
        if t:
            ultimo = t
    print('IMAGENS extraidas para %s:' % destino)
    for nome, ctx in achadas:
        print('   %s   <- depois de: %s' % (nome, ctx[:70]))
    print('Referencie no .md como: ![descricao](images/<arquivo>)')
    return achadas


# ---------------- normalização p/ casar com o .md ----------------
def strip_md(s):
    s = re.sub(r'<a id="[^"]*"></a>', '', s)
    s = re.sub(r'\[([^\]]+)\]\(#[^)]+\)', r'\1', s)
    s = s.replace('<br>', '').replace('**', '').replace('`', '')
    return s.strip()


def find_md_line(md_lines, texto):
    """Índices das linhas do .md cujo texto normalizado contém `texto`."""
    alvo = texto.strip()
    if len(alvo) < 4:
        return []
    return [i for i, l in enumerate(md_lines) if alvo in strip_md(l)]


# ---------------- principal ----------------
def main():
    if len(sys.argv) < 3:
        print(__doc__)
        return 1
    docx_path = os.path.abspath(sys.argv[1])
    md_path = os.path.abspath(sys.argv[2])
    aplicar = '--apply' in sys.argv

    for p in (docx_path, md_path):
        if not os.path.exists(p):
            print('ERRO: arquivo nao encontrado:', p)
            return 1

    if '--imagens' in sys.argv:
        extrair_imagens(docx_path, md_path)
        if '--apply' not in sys.argv:
            return 0

    # 1) .docx esperado a partir do .md atual
    tmp = os.path.join(tempfile.gettempdir(), '_esperado_sync.docx')
    md2docx.build(md_path, tmp)

    esperado = extract_blocks(tmp)
    editado = extract_blocks(docx_path)
    md_lines = open(md_path, encoding='utf-8').read().split('\n')

    sm = difflib.SequenceMatcher(a=esperado, b=editado, autojunk=False)
    ratio = sm.ratio()

    n_img_esp, n_img_edt = count_images(tmp), count_images(docx_path)

    print('=' * 70)
    print('SYNC  docx -> md')
    print('  docx :', docx_path)
    print('  md   :', md_path)
    print('  similaridade: %.4f | blocos: esperado=%d editado=%d'
          % (ratio, len(esperado), len(editado)))
    if n_img_edt != n_img_esp:
        print('  IMAGENS: esperado=%d  editado=%d  -> o humano inseriu/removeu imagem;'
              % (n_img_esp, n_img_edt))
        print('           extraia com --imagens e referencie no .md como ![alt](images/x.png)')
    print('=' * 70)

    from collections import Counter

    aplicadas = 0
    removidos, adicionados = [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            continue
        antigos = esperado[i1:i2]
        novos = editado[j1:j2]

        # caso seguro: substituição 1:1 de texto puro, ocorrência única no .md
        if tag == 'replace' and len(antigos) == len(novos):
            planos, todos_ok = [], True
            for a, b in zip(antigos, novos):
                linhas = find_md_line(md_lines, a)
                if len(linhas) == 1 and a in md_lines[linhas[0]]:
                    planos.append((linhas[0], a, b))
                else:
                    todos_ok = False
                    break
            if todos_ok:
                for idx, a, b in planos:
                    print('  [auto] md:%d  "%s"' % (idx + 1, a[:60]))
                    print('            -> "%s"' % b[:60])
                    if aplicar:
                        md_lines[idx] = md_lines[idx].replace(a, b, 1)
                aplicadas += len(planos)
                continue

        removidos.extend(antigos)
        adicionados.extend(novos)

    # Cancela ruído de alinhamento: textos que "somem" e "aparecem" na mesma
    # quantidade (cabeçalhos repetidos como CÓDIGO/DESCRIÇÃO) não são mudanças.
    comuns = Counter(removidos) & Counter(adicionados)
    def sem_ruido(itens):
        saldo = Counter(comuns)
        out = []
        for x in itens:
            if saldo.get(x, 0) > 0:
                saldo[x] -= 1
            else:
                out.append(x)
        return out
    removidos = sem_ruido(removidos)
    adicionados = sem_ruido(adicionados)
    # Se o texto "adicionado" já existe no .md, não é adição do humano — é ruído.
    adicionados = [b for b in adicionados if not find_md_line(md_lines, b)]
    # Células puramente numéricas (coluna ID dos quadros) são estrutura, não conteúdo.
    def relevante(x):
        return not re.fullmatch(r'[\d\W]{1,3}', x.strip())
    removidos = [a for a in removidos if relevante(a)]
    adicionados = [b for b in adicionados if relevante(b)]
    manuais = len(removidos) + len(adicionados)

    for a in removidos:
        locs = find_md_line(md_lines, a)
        loc = ('md:%d' % (locs[0] + 1)) if len(locs) == 1 else ('%d ocorr.' % len(locs))
        print('  [REMOVIDO no docx] (%s) %s' % (loc, a[:110]))
    for b in adicionados:
        print('  [ADICIONADO no docx]    %s' % b[:110])

    print('-' * 70)
    print('substituicoes seguras: %d %s | itens para decidir: %d'
          % (aplicadas, '(APLICADAS)' if aplicar else '(use --apply)', manuais))
    if aplicar and aplicadas:
        open(md_path, 'w', encoding='utf-8', newline='').write('\n'.join(md_lines))
        print('.md ATUALIZADO:', md_path)
    if manuais:
        print('Aja apenas nos blocos acima; nao releia o documento inteiro.')
    if aplicadas == 0 and manuais == 0:
        print('Nenhuma diferenca: .md e .docx estao sincronizados.')
    try:
        os.remove(tmp)
    except OSError:
        pass
    return 0


if __name__ == '__main__':
    sys.exit(main())
