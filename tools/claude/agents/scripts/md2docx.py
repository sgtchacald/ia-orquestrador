# -*- coding: utf-8 -*-
"""md2docx — converte um documento de análise .md (padrão 72b) em .docx formatado.

Uso:
    python md2docx.py <origem.md> <destino.docx>

O .md é a FONTE DA VERDADE; o .docx é um clone gerado a partir dele.
Genérico: serve para qualquer documento que siga o template 72b.

Suporta: tabelas markdown (com <br> nas células), QUADRO_DESCRITIVO de banco
(+ DDL em Courier), tabela de Endpoints, blocos ```sql```, cenários BDD,
imagens ![alt](caminho/relativo.png), **negrito**, `código`,
âncoras <a id="x"></a> -> bookmark e [rótulo](#x) -> hyperlink interno.
"""
import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = None   # definido em main() a partir de sys.argv
DEST = None

DARK   = '1F3864'
MED    = '2E75B6'
ZEBRA  = 'EBF3FB'
WHITE  = 'FFFFFF'
YELLOW = 'FFF2CC'
GRAY   = '444444'
FONT   = 'Calibri'

# ---------------- helpers ----------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'BFBFBF')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def new_table(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    set_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.allow_autofit = False
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    return t

USABLE_CM = 18.0   # A4 (21 cm) - margens laterais de 1,5 cm

def set_widths(t, cms):
    """Aplica larguras SEMPRE normalizadas para caber na área útil.

    Sem isso as tabelas estouram a margem e a última coluna sai cortada à
    direita — foi o que aconteceu no documento 01.
    """
    total = float(sum(cms)) or 1.0
    cms = [w * USABLE_CM / total for w in cms]
    for row in t.rows:
        cells = row.cells
        for i, w in enumerate(cms):
            if i < len(cells):
                cells[i].width = Cm(w)

_BM = [0]
_INLINE_RE = re.compile(r'(<a id="[^"]+"></a>|\[[^\]]+\]\(#[^)]+\))')

def _emit_plain(p, line, size, base_bold, color, base_font):
    idx = 0; bold = base_bold; code = False; buf = ''; segs = []
    while idx < len(line):
        if line[idx:idx+2] == '**':
            if buf:
                segs.append((buf, bold, code)); buf = ''
            bold = not bold; idx += 2; continue
        if line[idx] == '`':
            if buf:
                segs.append((buf, bold, code)); buf = ''
            code = not code; idx += 1; continue
        buf += line[idx]; idx += 1
    if buf:
        segs.append((buf, bold, code))
    for txt, bb, cc in segs:
        r = p.add_run(txt)
        r.bold = bb
        r.font.size = Pt(size)
        r.font.name = 'Courier New' if cc else base_font
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

def _emit_bookmark(p, name):
    _BM[0] += 1; bid = str(_BM[0])
    bs = OxmlElement('w:bookmarkStart'); bs.set(qn('w:id'), bid); bs.set(qn('w:name'), name)
    be = OxmlElement('w:bookmarkEnd'); be.set(qn('w:id'), bid)
    p._p.append(bs); p._p.append(be)

def _emit_hyperlink(p, anchor, label, size, base_font):
    hyper = OxmlElement('w:hyperlink'); hyper.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), base_font); rf.set(qn('w:hAnsi'), base_font); rPr.append(rf)
    c = OxmlElement('w:color'); c.set(qn('w:val'), '2E75B6'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    r.append(rPr)
    tn = OxmlElement('w:t'); tn.set(qn('xml:space'), 'preserve'); tn.text = label; r.append(tn)
    hyper.append(r); p._p.append(hyper)

def add_runs(p, text, size, base_bold=False, color=None, base_font=FONT):
    """Render text into paragraph p, interpreting \n (line break),
    **bold** / `code` toggles, <a id="X"></a> bookmarks and
    [label](#anchor) internal hyperlinks."""
    lines = text.split('\n')
    for li, line in enumerate(lines):
        if li > 0:
            p.add_run().add_break()
        for part in _INLINE_RE.split(line):
            if not part:
                continue
            ma = re.fullmatch(r'<a id="([^"]+)"></a>', part)
            if ma:
                _emit_bookmark(p, ma.group(1)); continue
            ml = re.fullmatch(r'\[([^\]]+)\]\(#([^)]+)\)', part)
            if ml:
                _emit_hyperlink(p, ml.group(2), ml.group(1), size, base_font); continue
            _emit_plain(p, part, size, base_bold, color, base_font)

def style_cell(cell, text, bold=False, white=False, bg=None, size=9,
               courier=False, align=None):
    if bg:
        set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    if align:
        p.alignment = align
    color = 'FFFFFF' if white else None
    base_font = 'Courier New' if courier else FONT
    add_runs(p, text, size, base_bold=bold, color=color, base_font=base_font)

# ---------------- paragraph styles ----------------
def _para(doc, text, size, bold=False, color=None, indent=0.0,
          before=2, after=2):
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    add_runs(p, text, size, base_bold=bold, color=color)
    return p

def h_title(doc, text):
    _para(doc, text, 16, bold=True, color=DARK, before=4, after=2)

def h_subtitle(doc, text):
    _para(doc, text, 13, bold=True, color=MED, before=0, after=6)

def h_meta(doc, text):
    _para(doc, text, 12, bold=True, color=DARK, before=10, after=4)

def h_section(doc, text):
    _para(doc, text, 14, bold=True, color=DARK, before=12, after=4)

def h_sub(doc, text):
    _para(doc, text, 12, bold=True, color=DARK, before=8, after=3)

def body(doc, text):
    _para(doc, text, 10, before=2, after=2)

def bullet(doc, text):
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    add_runs(p, '\u2022  ' + text, 10)

def quote(doc, text):
    _para(doc, text, 9, color=GRAY, before=4, after=2)

def bdd_line(doc, text):
    m = re.match(r'^(Dado que|Dado|Quando|Então|E que|E)\b', text)
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    if m:
        kw = m.group(1)
        rest = text[len(kw):]
        r = p.add_run(kw); r.bold = True; r.font.size = Pt(10); r.font.name = FONT
        add_runs(p, rest, 10)
    else:
        add_runs(p, text, 10)

def code_block(doc, code, size=8):
    p = doc.add_paragraph(style='Normal')
    pf = p.paragraph_format
    pf.left_indent = Cm(0.3)
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    lines = code.split('\n')
    for i, l in enumerate(lines):
        if i > 0:
            br = p.add_run(); br.add_break()
        r = p.add_run(l if l != '' else '')
        r.font.name = 'Courier New'
        r.font.size = Pt(size)

# ---------------- table renderers ----------------
def std_table(doc, rows, widths=None):
    headers = rows[0]
    data = rows[1:]
    cols = len(headers)
    t = new_table(doc, 1 + len(data), cols)
    for j, h in enumerate(headers):
        style_cell(t.rows[0].cells[j], h, bold=True, white=True, bg=DARK)
    for i, row in enumerate(data):
        bg = WHITE if i % 2 == 0 else ZEBRA
        for j in range(cols):
            val = row[j] if j < len(row) else ''
            style_cell(t.rows[i+1].cells[j], val, bg=bg)
    # sem largura definida -> distribui a área útil igualmente (evita estouro)
    set_widths(t, widths if widths else [1.0] * cols)
    return t

def endpoints_table(doc, rows):
    header = rows[0]
    body_rows = rows[1:]
    cols = len(header)  # 5
    # build sequentially
    t = new_table(doc, 1, cols)
    for j, h in enumerate(header):
        style_cell(t.rows[0].cells[j], h, bold=True, white=True, bg=DARK)
    for row in body_rows:
        first = row[0].strip()
        nonempty = [c for c in row if c.strip()]
        is_data = first.upper().startswith('EDP') and len(nonempty) >= 2
        r = t.add_row()
        if is_data:
            for j in range(cols):
                val = row[j] if j < len(row) else ''
                style_cell(r.cells[j], val, bg=WHITE)
        else:
            m = r.cells[0].merge(r.cells[cols-1])
            style_cell(m, row[0], bg=ZEBRA)
    set_widths(t, [1.6, 1.3, 3.2, 7.4, 2.0])
    return t

_QUADRO_COUNT = [0]

def quebra_pagina(doc):
    """Separa visualmente um QUADRO_DESCRITIVO do anterior."""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

def banco_quadro(doc, rows, ddl_code):
    # a partir do segundo quadro, começa em página nova (pedido do usuário:
    # "separe os quadros descritivos")
    _QUADRO_COUNT[0] += 1
    if _QUADRO_COUNT[0] > 1:
        quebra_pagina(doc)
    # locate ALTERAÇÃO row
    idx_alter = None
    for i, r in enumerate(rows):
        if 'ALTERAÇÃO' in r[0].upper():
            idx_alter = i
            break
    name = rows[0][0]
    tabela = rows[1][0].replace('**', '').strip()
    obs = rows[2][0]
    header = rows[3]
    fields = rows[4:idx_alter]
    cols = 4
    total = 3 + 1 + len(fields) + 1 + 1
    t = new_table(doc, total, cols)
    # r0 name
    m = t.rows[0].cells[0].merge(t.rows[0].cells[cols-1])
    style_cell(m, name, bold=True, white=True, bg=DARK)
    # r1 tabela
    m = t.rows[1].cells[0].merge(t.rows[1].cells[cols-1])
    style_cell(m, 'TABELA DO BANCO DE DADOS: ' + tabela.replace('TABELA DO BANCO DE DADOS: ', ''),
               bold=True, white=True, bg=MED)
    # r2 obs
    m = t.rows[2].cells[0].merge(t.rows[2].cells[cols-1])
    style_cell(m, obs, bg=YELLOW, size=9)
    # r3 header
    for j, h in enumerate(header):
        style_cell(t.rows[3].cells[j], h, bold=True, white=True, bg=DARK)
    # fields
    base = 4
    for i, f in enumerate(fields):
        bg = WHITE if i % 2 == 0 else ZEBRA
        for j in range(cols):
            val = f[j] if j < len(f) else ''
            style_cell(t.rows[base+i].cells[j], val, bg=bg)
    # alteração row
    r_alter = base + len(fields)
    m = t.rows[r_alter].cells[0].merge(t.rows[r_alter].cells[cols-1])
    style_cell(m, 'ALTERAÇÃO NA ESTRUTURA DO BANCO DE DADOS', bold=True, white=True, bg=MED)
    # ddl row
    r_ddl = r_alter + 1
    m = t.rows[r_ddl].cells[0].merge(t.rows[r_ddl].cells[cols-1])
    style_cell(m, ddl_code, bg=WHITE, size=8, courier=True)
    set_widths(t, [1.0, 3.6, 6.9, 4.5])
    return t

# ---------------- tokenizer ----------------
def tokenize(md):
    lines = md.split('\n')
    toks = []
    i = 0
    n = len(lines)
    def is_sep(cells):
        return all(re.fullmatch(r':?-{2,}:?', c.strip() or '-') or set(c.strip()) <= set('-:') and c.strip() != '' for c in cells) if cells else False
    while i < n:
        line = lines[i]
        s = line.strip()
        if s.startswith('```'):
            lang = s[3:].strip()
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            code = '\n'.join(buf)
            toks.append(('sql' if lang == 'sql' else 'code', code))
            continue
        if s.startswith('|'):
            tbl = []
            while i < n and lines[i].strip().startswith('|'):
                raw = lines[i].strip()
                parts = raw.split('|')[1:-1]
                cells = [c.strip().replace('<br>', '\n') for c in parts]
                # detect separator
                joined = ''.join(c.replace(' ', '') for c in cells)
                if joined and set(joined) <= set('-:'):
                    i += 1
                    continue
                tbl.append(cells)
                i += 1
            if tbl:
                toks.append(('table', tbl))
            continue
        if s == '':
            i += 1
            continue
        if s == '---':
            toks.append(('hr',))
            i += 1
            continue
        if s.startswith('### '):
            toks.append(('h', 3, s[4:].strip()))
            i += 1
            continue
        if s.startswith('## '):
            toks.append(('h', 2, s[3:].strip()))
            i += 1
            continue
        if s.startswith('# '):
            toks.append(('h', 1, s[2:].strip()))
            i += 1
            continue
        mimg = re.fullmatch(r'!\[([^\]]*)\]\(([^)]+)\)', s)
        if mimg:
            toks.append(('image', mimg.group(2), mimg.group(1)))
            i += 1
            continue
        if s.startswith('> '):
            toks.append(('quote', s[2:].strip()))
            i += 1
            continue
        if s.startswith('- '):
            toks.append(('bullet', s[2:].strip()))
            i += 1
            continue
        toks.append(('body', s))
        i += 1
    return toks

# ---------------- render ----------------
def is_banco_quadro(rows):
    return rows and rows[0] and rows[0][0].strip().upper().startswith('QUADRO_DESCRITIVO')

def is_endpoints(rows):
    if not rows:
        return False
    h = [c.strip().upper() for c in rows[0]]
    return h[:5] == ['CÓDIGO', 'HTTP', 'PERMISSÃO', 'PATH', 'FINALIZADO?']

def render(doc, toks):
    from collections import deque
    dq = deque(toks)
    in_bdd = False
    while dq:
        tok = dq.popleft()
        kind = tok[0]
        if kind == 'h':
            lvl, txt = tok[1], tok[2]
            if lvl == 1:
                h_title(doc, txt)
            elif lvl == 2:
                if txt.startswith('Módulo'):
                    h_subtitle(doc, txt)
                elif re.match(r'^\d+\.', txt):
                    in_bdd = txt.startswith('16.')
                    h_section(doc, txt)
                else:
                    h_meta(doc, txt)
            else:  # lvl 3
                if txt.upper().startswith('QUADRO_DESCRITIVO'):
                    continue  # bare quadro label, skip (name shown in merged row)
                h_sub(doc, txt)
        elif kind == 'hr':
            continue
        elif kind == 'body':
            if in_bdd:
                bdd_line(doc, tok[1])
            else:
                body(doc, tok[1])
        elif kind == 'bullet':
            bullet(doc, tok[1])
        elif kind == 'image':
            import os
            rel = tok[1].replace('/', os.sep)
            path = rel if os.path.isabs(rel) else os.path.join(os.path.dirname(SRC), rel)
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(4); pf.space_after = Pt(6)
            if os.path.exists(path):
                p.add_run().add_picture(path, width=Cm(16.0))
            else:
                r = p.add_run('[imagem não encontrada: %s]' % tok[1])
                r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GRAY)
        elif kind == 'quote':
            quote(doc, tok[1])
        elif kind == 'sql':
            code_block(doc, tok[1], size=8)
        elif kind == 'code':
            code_block(doc, tok[1], size=8)
        elif kind == 'table':
            rows = tok[1]
            if is_banco_quadro(rows):
                # consume following sql token as DDL
                ddl = ''
                if dq and dq[0][0] == 'sql':
                    ddl = dq.popleft()[1]
                banco_quadro(doc, rows, ddl)
            elif is_endpoints(rows):
                endpoints_table(doc, rows)
            else:
                # choose widths by header signature
                header = [c.strip().upper() for c in rows[0]]
                widths = None
                if header[:4] == ['ID', 'DESCRIÇÃO', 'PRIORIDADE', 'SITUAÇÃO']:
                    widths = [1.2, 11.6, 2.0, 2.2]
                elif header[:4] == ['ID', 'CATEGORIA', 'DESCRIÇÃO', 'CRITÉRIO DE ACEITAÇÃO']:
                    widths = [1.3, 2.3, 7.0, 6.4]
                elif header[:4] == ['CÓDIGO', 'NOME', 'ATOR PRINCIPAL', 'DESCRIÇÃO']:
                    widths = [1.6, 3.2, 2.6, 9.6]
                elif header[:4] == ['ID', 'NOME', 'PROPRIEDADES', 'OBSERVAÇÕES']:
                    widths = [1.0, 3.4, 6.7, 5.9]
                elif header[:3] == ['ID', 'NOME', 'DESCRIÇÃO']:
                    widths = [1.6, 3.0, 12.4]
                elif header[:2] == ['ID', 'DESCRIÇÃO']:
                    widths = [1.6, 15.4]
                elif header[:2] == ['CÓDIGO', 'DESCRIÇÃO']:
                    widths = [1.9, 15.1]
                elif header[:3] == ['CÓDIGO', 'NOME', 'DESCRIÇÃO']:
                    widths = [1.9, 3.0, 12.1]
                elif header[:3] == ['CÓDIGO', 'DESCRIÇÃO', 'PERFIS COM ACESSO']:
                    widths = [1.9, 11.5, 3.6]
                elif header[:3] == ['PARÂMETRO', 'VALOR PADRÃO', 'DESCRIÇÃO']:
                    widths = [5.5, 2.6, 8.9]
                elif header[:3] == ['TABELA PRINCIPAL', 'TABELA DE AUDITORIA', 'CAMPOS AUDITADOS']:
                    widths = [3.5, 3.5, 10.0]
                elif header[:3] == ['Nº', 'OBSERVAÇÃO', 'REFERÊNCIA / IMPACTO']:
                    widths = [1.0, 12.4, 3.6]
                std_table(doc, rows, widths)

# ---------------- main ----------------
def build(src_md, dest_docx):
    """Gera dest_docx a partir de src_md. Reutilizável por outros scripts."""
    global SRC, DEST
    SRC, DEST = src_md, dest_docx
    return main()

def main():
    with open(SRC, encoding='utf-8') as f:
        md = f.read()
    toks = tokenize(md)
    doc = Document()
    # margins
    # Página A4 explícita: o padrão do python-docx é Letter (21,59 cm), o que
    # dava área útil de 16,6 cm com margens de 2,5 cm — as tabelas de 17 cm
    # estouravam e a última coluna saía cortada à direita.
    # A4 (21 cm) com margens de 1,5 cm => área útil de exatamente USABLE_CM.
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
    # base style font
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(10)
    render(doc, toks)
    doc.save(DEST)
    print('SAVED:', DEST)
    return DEST

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    SRC, DEST = os.path.abspath(sys.argv[1]), os.path.abspath(sys.argv[2])
    if not os.path.exists(SRC):
        print('ERRO: origem .md nao encontrada:', SRC)
        sys.exit(1)
    # Guarda de segurança: nao sobrescrever um .docx editado depois do .md
    if os.path.exists(DEST) and '--force' not in sys.argv:
        if os.path.getmtime(DEST) > os.path.getmtime(SRC) + 5:
            print('ABORTADO: o .docx e MAIS NOVO que o .md.')
            print('  O humano provavelmente editou o Word. Rode antes:')
            print('    python sync_docx_md.py "%s" "%s"' % (DEST, SRC))
            print('  (ou use --force para sobrescrever mesmo assim)')
            sys.exit(2)
    main()
